from flask import Flask, Response, render_template, request, redirect, url_for, flash, jsonify
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from flask_migrate import Migrate
from models import db, Document, Title, Responsibility, PubAuthority, PubPlace, PubDate, Identifier, License, Source, Note, Description, Abstract, CreationPlace, CreationDate, Language, Category
from download_tei import get_data, generate_tei_header, generate_tei_content, download_all_documents_as_tei_zip
import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument



app = Flask(__name__, static_url_path='/static')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///test2.db'
app.secret_key = 'mysecretkey'
db.init_app(app)
migrate = Migrate(app, db)


@app.route('/')
def index():
    documents = Document.query.order_by(Document.updated_at.desc()).all()
    return render_template('index.html', documents=documents)


@app.route('/document/new', methods=['GET', 'POST'])
def new_document():
    if request.method == 'POST':
        title = request.form['docTitle']
        content = request.form['content']
        if not title or not content:
            flash('Please enter a title and content for your document.', 'error')
            return redirect(url_for('new_document'))
        document = Document(docTitle=title, content=content)
        db.session.add(document)
        db.session.commit()
        flash('Your document has been created successfully.', 'success')
        return redirect(url_for('index'))
    return render_template('new_document.html')


ALLOWED_EXTENSIONS = {'docx', 'pdf', 'txt'}

# Helper function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Helper function to read text content from a DOCX file
def read_docx(uploaded_file):
    try:
        doc = DocxDocument(uploaded_file)
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        return text_content
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return None

# Helper function to read text content from a PDF file
def read_pdf(uploaded_file):
    try:
        pdf = PdfReader(uploaded_file)
        text_content = ""
        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            text_content += page.extract_text() + "\n"
        return text_content
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        return None

# Helper function to read text content from a TXT file
def read_txt(uploaded_file):
    try:
        text_content = uploaded_file.read().decode('utf-8')
        return text_content
    except Exception as e:
        print(f"Error reading TXT file: {e}")
        return None

# Main route for uploading documents
@app.route('/document/upload', methods=['GET', 'POST'])
def upload_document():
    if request.method == 'POST':
        title = request.form['docTitle']
        uploaded_file = request.files['uploaded_file']

        if not title:
            flash('Please enter a title for your document.', 'error')
        elif not uploaded_file:
            flash('Please select a file to upload.', 'error')
        elif not allowed_file(uploaded_file.filename):
            flash('Invalid file format. Please upload a supported file type.', 'error')
        else:
            # Determine the file extension and call the appropriate function
            file_extension = uploaded_file.filename.rsplit('.', 1)[1].lower()
            if file_extension == 'docx':
                text_content = read_docx(uploaded_file)
            elif file_extension == 'pdf':
                text_content = read_pdf(uploaded_file)
            elif file_extension == 'txt':
                text_content = read_txt(uploaded_file)
            else:
                text_content = None
            
            if text_content is not None:
                # Create a new Document object and save it to the database
                document = Document(docTitle=title, content=text_content)
                db.session.add(document)
                db.session.commit()
                flash('Your document has been uploaded successfully.', 'success')
                return redirect(url_for('index'))

    return render_template('new_document.html')


@app.route('/document/<int:id>/edit', methods=['GET', 'POST'])
def edit_document(id):
    document = Document.query.get_or_404(id)
    return render_template('edit_document.html', document=document)


# Flask route to save or update a document
@app.route("/save_document", methods=["POST"])
def save_document():
    data = request.get_json()
    document_id = data.get("document_id")
    title = data.get("docTitle")
    content = data.get("content")

    # Check if a document with the given ID already exists
    document = Document.query.filter_by(id=document_id).first()

    if document:
        # Update the existing document
        document.docTitle = title
        document.content = content
        document.updated_at = datetime.utcnow()
    else:
        # Create a new document
        new_document = Document(docTitle=title, content=content)
        db.session.add(new_document)

    db.session.commit()
    return jsonify({"message": "Document saved successfully!"})


# Flask route to fetch a document by its ID
@app.route("/get_document/<int:document_id>")
def get_document(document_id):
    document = Document.query.get(document_id)
    if document:
        return jsonify({"docTitle": document.docTitle, "content": document.content})
    else:
        return jsonify({"error": "Document not found"}), 404


@app.route('/save_metadata/<int:id>', methods=['POST'])
def save_metadata(id):
    document = Document.query.get_or_404(id)
    
    if request.method == 'POST':

        title_texts = request.form.getlist('title-text')
        title_languages = request.form.getlist('title-language')
        existing_titles = Title.query.filter_by(document_id=document.id).all()
        # Create dictionaries to track processed titles
        processed_titles = {}
        # Update existing titles
        for i, existing_title in enumerate(existing_titles):
            if i < len(title_texts):
                existing_title.text = title_texts[i]
                if i < len(title_languages):
                    existing_title.language = title_languages[i]
                else:
                    existing_title.language = None
            processed_titles[i] = True
        # Add new titles
        for i, text in enumerate(title_texts):
            if i not in processed_titles:
                new_title = Title(document_id=document.id, text=text)
                if i < len(title_languages):
                    new_title.language = title_languages[i]
                else:
                    new_title.language = None
                db.session.add(new_title)

        resp_surnames = request.form.getlist('responsibility-surname')
        resp_names = request.form.getlist('responsibility-name')
        resp_authorities = request.form.getlist('responsibility-authority')
        resp_roles = request.form.getlist('responsibility-role')
        existing_responsibilities = Responsibility.query.filter_by(document_id=document.id).all()
        # Create dictionaries to track processed resps
        processed_responsibilities = {}
        # Update existing titles
        for i, existing_responsibility in enumerate(existing_responsibilities):
            if i < len(resp_surnames):
                existing_responsibility.surname = resp_surnames[i]
                if i < len(resp_names):
                    existing_responsibility.name = resp_names[i]
                if i < len(resp_authorities):
                    existing_responsibility.authority = resp_authorities[i]
                else:
                    existing_responsibility.authority = None
                if i < len(resp_roles):
                    existing_responsibility.role = resp_roles[i]
            processed_responsibilities[i] = True
        # Add new resps
        for i, surname in enumerate(resp_surnames):
            if i not in processed_responsibilities:
                new_responsibility = Responsibility(document_id=document.id, 
                                          surname=surname)
                if i < len(resp_names):
                    new_responsibility.name = resp_names[i]
                else:
                    new_responsibility.name = None
                if i < len(resp_authorities):
                    new_responsibility.authority = resp_authorities[i]
                else:
                    new_responsibility.authority = None
                if i < len(resp_roles):
                    new_responsibility.role = resp_roles[i]
                else:
                    new_responsibility.role = None
                db.session.add(new_responsibility)

        pubAuthority_names = request.form.getlist('pubAuthority-name')
        pubAuthority_authorities = request.form.getlist('pubAuthority-authority')
        pubAuthority_roles = request.form.getlist('pubAuthority-role')
        existing_pubAuthorities = PubAuthority.query.filter_by(document_id=document.id).all()
        # Create dictionaries to track processed pubauths
        processed_pubAuthorities = {}
        # Update existing pubauths
        for i, existing_pubAuthority in enumerate(existing_pubAuthorities):
            if i < len(pubAuthority_names):
                existing_pubAuthority.name = pubAuthority_names[i]
                if i < len(pubAuthority_authorities):
                    existing_pubAuthority.authority = pubAuthority_authorities[i]
                else:
                    existing_pubAuthority.authority = None
                if i < len(resp_roles):
                    existing_pubAuthority.role = pubAuthority_roles[i]
            processed_pubAuthorities[i] = True
        # Add new pubauths
        for i, name in enumerate(pubAuthority_names):
            if i not in processed_pubAuthorities:
                new_pubAuthority = PubAuthority(document_id=document.id, 
                                                name=name)
                if i < len(pubAuthority_authorities):
                    new_pubAuthority.authority = pubAuthority_authorities[i]
                else:
                    new_pubAuthority.authority = None
                if i < len(pubAuthority_roles):
                    new_pubAuthority.role = pubAuthority_roles[i]
                else:
                    new_pubAuthority.role = None
                db.session.add(new_pubAuthority)

        pubPlace_name = request.form.get('pubPlace-name')
        pubPlace_authority = request.form.get('pubPlace-authority')
        existing_pubPlace = PubPlace.query.filter_by(document_id=document.id).first()
        if existing_pubPlace:
            existing_pubPlace.name = pubPlace_name
            existing_pubPlace.authority = pubPlace_authority
        else:
            pubPlace = PubPlace(document_id=document.id,
                                name=pubPlace_name,
                                authority=pubPlace_authority)
            db.session.add(pubPlace)

        pubDate_date = datetime.strptime(request.form.get('pubDate-date'), '%Y-%m-%d').date()
        existing_pubDate = PubDate.query.filter_by(document_id=document.id).first()
        if existing_pubDate:
            existing_pubDate.date = pubDate_date
        else:
            pubDate = PubDate(document_id=document.id,
                              date=pubDate_date)
            db.session.add(pubDate)


        ident_texts = request.form.getlist('identifier-text')
        ident_types = request.form.getlist('identifier-type')
        existing_identifiers = Identifier.query.filter_by(document_id=document.id).all()
        # Create dictionaries to track processed ids
        processed_identifiers = {}
        # Update existing ids
        for i, existing_identifier in enumerate(existing_identifiers):
            if i < len(ident_texts):
                existing_identifier.text = ident_texts[i]
                if i < len(ident_types):
                    existing_identifier.type = ident_types[i]
                else:
                    existing_identifier.type = None
            processed_identifiers[i] = True
        # Add new ids
        for i, text in enumerate(ident_texts):
            if i not in processed_identifiers:
                new_ident = Identifier(document_id=document.id, 
                                       text=text)
                if i < len(ident_types):
                    new_ident.type = ident_types[i]
                else:
                    new_ident.type = None
                db.session.add(new_ident)


        license_text = request.form.get('license-text')
        license_link = request.form.get('license-link')
        existing_license = License.query.filter_by(document_id=document.id).first()
        if existing_license:
            existing_license.text = license_text
            existing_license.link = license_link
        else:
            license = License(document_id=document.id,
                                        text=license_text,
                                        link=license_link)
            db.session.add(license)

        source_text = request.form.get('source-text')
        existing_source = Source.query.filter_by(document_id=document.id).first()
        if existing_source:
            existing_source.text = source_text
        else:
            source = Source(document_id=document.id,
                            text=source_text)
            db.session.add(source)

        note_text = request.form.get('note-text')
        existing_note = Note.query.filter_by(document_id=document.id).first()
        if existing_note:
            existing_note.text = note_text
        else:
            note = Note(document_id=document.id,
                        text=note_text)
            db.session.add(note)

        description_text = request.form.get('description-text')
        existing_description = Description.query.filter_by(document_id=document.id).first()
        if existing_description:
            existing_description.text = description_text
        else:
            description = Description(document_id=document.id,
                                      text=description_text)
            db.session.add(description)
        
        abstract_text = request.form.get('abstract-text')
        existing_abstract = Abstract.query.filter_by(document_id=document.id).first()
        if existing_abstract:
            existing_abstract.text = abstract_text
        else:
            abstract = Abstract(document_id=document.id,
                                text=abstract_text)
            db.session.add(abstract)

        creationPlace_name = request.form.get('creationPlace-name')
        creationPlace_authority = request.form.get('creationPlace-authority')
        existing_creationPlace = CreationPlace.query.filter_by(document_id=document.id).first()
        if existing_creationPlace:
            existing_creationPlace.name = creationPlace_name
            existing_creationPlace.authority = creationPlace_authority
        else:
            creationPlace = CreationPlace(document_id=document.id,
                                name=creationPlace_name,
                                authority=creationPlace_authority)
            db.session.add(creationPlace)

        creationDate_date = datetime.strptime(request.form.get('creationDate-date'), '%Y-%m-%d').date()
        existing_creationDate = CreationDate.query.filter_by(document_id=document.id).first()
        if existing_creationDate:
            existing_creationDate.date = creationDate_date
        else:
            creationDate = CreationDate(document_id=document.id,
                              date=creationDate_date)
            db.session.add(creationDate)

        lang_text = request.form.get('language-text')
        lang_ident = request.form.get('language-ident')
        lang_usage = request.form.get('language-usage')
        existing_lang = Language.query.filter_by(document_id=document.id).first()
        if existing_lang:
            existing_lang.text = lang_text
            existing_lang.ident = lang_ident
            existing_lang.usage = lang_usage
        else:
            lang = Language(document_id=document.id,
                            text=lang_text,
                            ident=lang_ident,
                            usage=lang_usage)
            db.session.add(lang)

        category_type = request.form.get('category-type')
        existing_category = Category.query.filter_by(document_id=document.id).first()
        if existing_category:
            existing_category.type = category_type
        else:
            category = Category(document_id=document.id,
                                type=category_type)
            db.session.add(category)

        db.session.commit()

    return redirect(url_for('edit_document', id=document.id))


@app.route('/get_existing_data/<int:id>')
def get_existing_data(id):
    existing_data = {
        'title': [],
        'responsibility': [],
        'pubAuthority': [],
        'pubPlace': [],
        'pubDate': [],
        'identifier': [],
        'license': [],
        'source': [],
        'note': [],
        'desc': [],
        'abstract': [],
        'creationPlace': [],
        'creationDate': [],
        'lang': [],
        'category': [],
        # Add more field types as needed
    }

    existing_titles = Title.query.filter_by(document_id=id).all()
    for title in existing_titles:
        existing_data['title'].append({
            'text': title.text,
            'language': title.language,
        })
        print(title.text, title.language)
    
    existing_resps = Responsibility.query.filter_by(document_id=id).all()
    for resp in existing_resps:
        existing_data['responsibility'].append({
            'surname': resp.surname,
            'name': resp.name,
            'authority': resp.authority,
            'role': resp.role
        })

    existing_pubAuthorities = PubAuthority.query.filter_by(document_id=id).all()
    for pubAuthority in existing_pubAuthorities:
        existing_data['pubAuthority'].append({
            'name': pubAuthority.name,
            'authority': pubAuthority.authority,
            'role': pubAuthority.role
        })
    
    existing_pubPlace = PubPlace.query.filter_by(document_id=id).all()
    for pubPlace in existing_pubPlace:
        existing_data['pubPlace'].append({
            'name': pubPlace.name,
            'authority': pubPlace.authority
        })
    
    existing_pubDate = PubDate.query.filter_by(document_id=id).all()
    for pubDate in existing_pubDate:
        existing_data['pubDate'].append({
            'date': pubDate.date
        })

    existing_identifiers = Identifier.query.filter_by(document_id=id).all()
    for ident in existing_identifiers:
        existing_data['identifier'].append({
            'text': ident.text,
            'type': ident.type
        })
    
    existing_license = License.query.filter_by(document_id=id).all()
    for license in existing_license:
        existing_data['license'].append({
            'text': license.text,
            'link': license.link
        })

    existing_sources = Source.query.filter_by(document_id=id).all()
    for source in existing_sources:
        existing_data['source'].append({
            'text': source.text
        })

    existing_notes = Note.query.filter_by(document_id=id).all()
    for note in existing_notes:
        existing_data['note'].append({
            'text': note.text
        })
    
    existing_description = Description.query.filter_by(document_id=id).all()
    for desc in existing_description:
        existing_data['desc'].append({
            'text': desc.text
        })
    
    existing_abstract = Abstract.query.filter_by(document_id=id).all()
    for abstract in existing_abstract:
        existing_data['abstract'].append({
            'text': abstract.text
        })

    existing_creationPlace = CreationPlace.query.filter_by(document_id=id).all()
    for creationPlace in existing_creationPlace:
        existing_data['creationPlace'].append({
            'name': creationPlace.name,
            'authority': creationPlace.authority
        })
    
    existing_creationDate = CreationDate.query.filter_by(document_id=id).all()
    for creationDate in existing_creationDate:
        existing_data['creationDate'].append({
            'date': creationDate.date
        })
    
    existing_languages = Language.query.filter_by(document_id=id).all()
    for lang in existing_languages:
        existing_data['lang'].append({
            'text': lang.text,
            'ident': lang.ident,
            'usage': lang.usage
        })
    
    existing_categories = Category.query.filter_by(document_id=id).all()
    for cat in existing_categories:
        existing_data['category'].append({
            'type': cat.type
        })

    return jsonify(existing_data)


@app.route('/delete/<model_name>/<int:record_id>', methods=['DELETE'])
def delete_record(model_name, record_id):
    try:
        # Get the SQLAlchemy model class based on the provided model_name
        model_class = globals().get(model_name)
        
        if not model_class:
            return jsonify({'message': 'Model not found'}), 404

        # Retrieve the record by its ID
        record = model_class.query.get(record_id)

        if record:
            # Delete the record from the database
            db.session.delete(record)
            db.session.commit()
            return jsonify({'message': f'{model_name} deleted successfully'}), 200
        else:
            return jsonify({'message': f'{model_name} not found'}), 404

    except Exception as e:
        return jsonify({'message': 'An error occurred'}), 500


@app.route('/download_tei/<int:document_id>')
def download_tei(document_id):
    
    data = get_data(document_id=document_id)
    tei_header = generate_tei_header(data=data)
    tei_content = generate_tei_content(data=data)

    # Create the complete TEI document
    tei_template = f"""
    <TEI xmlns="http://www.tei-c.org/ns/1.0">
      {tei_header}
      {tei_content}
    </TEI>
    """

    # Return the TEI document as a downloadable file
    response = Response(tei_template, content_type='application/xml')
    response.headers['Content-Disposition'] = f'attachment; filename=document_{document_id}.xml'

    return response


@app.route('/download_tei/all_documents')
def download_all_documents_route():
    try:
        # Call the download function from download.py
        zip_filename = download_all_documents_as_tei_zip()
        
        # Check if the file exists
        if not os.path.isfile(zip_filename):
            return "ZIP file not found", 404

        # Open the ZIP file in binary read mode
        with open(zip_filename, 'rb') as zip_file:
            zip_data = zip_file.read()

        # Set the appropriate headers for downloading a ZIP file
        response = Response(zip_data)
        response.headers['Content-Type'] = 'application/zip'
        response.headers['Content-Disposition'] = 'attachment; filename=tei_documents.zip'

        return response
    except Exception as e:
        return str(e), 500


@app.route('/docta')
def serve_docta():
    rdf_file_path = 'instance/docta.ttl'

    with open(rdf_file_path, 'rb') as f:
        rdf_content = f.read()

    response = app.response_class(
        response=rdf_content,
        status=200,
        mimetype='application/turtle'
    )

    return response


if __name__ == '__main__':
    app.run(debug=True)

